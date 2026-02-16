"""
File handler — reads/writes txt/doc/docx files and mapping JSON files.
Supports same-format output: input DOCX -> output DOCX, etc.
Processes headers, footers, tables, and document properties.
"""

import os
import io
import json
import subprocess
import tempfile
from datetime import datetime
from docx import Document

MAPPINGS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "mappings")


def read_uploaded_file(uploaded_file) -> str:
    """
    Read uploaded file content. Supports .txt / .doc / .docx.

    Args:
        uploaded_file: Streamlit UploadedFile object

    Returns:
        File content as string
    """
    filename = uploaded_file.name.lower()

    if filename.endswith(".docx"):
        return _read_docx(uploaded_file)
    elif filename.endswith(".doc"):
        return _read_doc(uploaded_file)
    else:
        return _read_txt(uploaded_file)


def get_uploaded_bytes(uploaded_file) -> bytes:
    """
    Get raw bytes from an uploaded file (resets read position afterward).

    Args:
        uploaded_file: Streamlit UploadedFile object

    Returns:
        Raw file bytes
    """
    uploaded_file.seek(0)
    data = uploaded_file.read()
    uploaded_file.seek(0)
    return data


def _read_txt(uploaded_file) -> str:
    """Read a .txt file with UTF-8/GBK fallback."""
    content = uploaded_file.read()
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("gbk")


def _read_docx(uploaded_file) -> str:
    """Read a .docx file, extracting all paragraph text."""
    content = uploaded_file.read()
    doc = Document(io.BytesIO(content))
    paragraphs = [para.text for para in doc.paragraphs]
    return "\n".join(paragraphs)


def _read_doc(uploaded_file) -> str:
    """Read a .doc file via macOS textutil conversion."""
    content = uploaded_file.read()

    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp_doc:
        tmp_doc.write(content)
        tmp_doc_path = tmp_doc.name

    tmp_txt_path = tmp_doc_path.replace(".doc", ".txt")

    try:
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-output", tmp_txt_path, tmp_doc_path],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode != 0:
            raise RuntimeError(f"textutil conversion failed: {result.stderr}")

        with open(tmp_txt_path, "r", encoding="utf-8") as f:
            return f.read()
    finally:
        for path in [tmp_doc_path, tmp_txt_path]:
            if os.path.exists(path):
                os.unlink(path)


# ============================================================
# Same-format output: apply replacements to DOCX/DOC files
# ============================================================

def apply_replacements_to_docx(docx_bytes: bytes, replacements: list[tuple[str, str]]) -> bytes:
    """
    Apply text replacements to a DOCX file while preserving formatting.
    Processes body paragraphs, tables, headers, footers, and core properties.

    Args:
        docx_bytes: Original DOCX file bytes
        replacements: List of (old_text, new_text) tuples, sorted by old_text length descending

    Returns:
        Modified DOCX as bytes
    """
    doc = Document(io.BytesIO(docx_bytes))

    def _replace_in_paragraph(para):
        runs = para.runs
        if not runs:
            return
        full_text = "".join(run.text for run in runs)
        new_text = full_text
        for old, new in replacements:
            new_text = new_text.replace(old, new)
        if new_text != full_text:
            runs[0].text = new_text
            for run in runs[1:]:
                run.text = ""

    def _replace_in_container(container):
        """Process all paragraphs and tables in a container (body, header, footer, cell)."""
        for para in container.paragraphs:
            _replace_in_paragraph(para)
        if hasattr(container, 'tables'):
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            _replace_in_paragraph(para)

    # Body paragraphs and tables
    _replace_in_container(doc)

    # Headers and footers (all sections)
    for section in doc.sections:
        for hf_attr in ['header', 'footer', 'first_page_header', 'first_page_footer',
                        'even_page_header', 'even_page_footer']:
            try:
                hf = getattr(section, hf_attr)
                _replace_in_container(hf)
            except Exception:
                continue

    # Core properties (title, author, subject, etc.)
    try:
        props = doc.core_properties
        for attr in ['title', 'subject', 'author', 'comments', 'description',
                     'last_modified_by', 'keywords', 'category']:
            try:
                val = getattr(props, attr, None)
                if val and isinstance(val, str):
                    new_val = val
                    for old, new in replacements:
                        new_val = new_val.replace(old, new)
                    if new_val != val:
                        setattr(props, attr, new_val)
            except Exception:
                continue
    except Exception:
        pass

    # Extended properties (company, manager) — modify raw XML
    try:
        for rel in doc.part.rels.values():
            if 'extended-properties' in str(getattr(rel, 'reltype', '')):
                app_part = rel.target_part
                xml_str = app_part.blob.decode('utf-8')
                new_xml = xml_str
                for old, new in replacements:
                    new_xml = new_xml.replace(old, new)
                if new_xml != xml_str:
                    app_part._blob = new_xml.encode('utf-8')
                break
    except Exception:
        pass

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def apply_replacements_to_doc(doc_bytes: bytes, replacements: list[tuple[str, str]]) -> bytes:
    """
    Apply text replacements to a .doc file.
    Converts .doc -> .docx, applies replacements, converts back to .doc.

    Args:
        doc_bytes: Original .doc file bytes
        replacements: List of (old_text, new_text) tuples

    Returns:
        Modified .doc as bytes
    """
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        tmp.write(doc_bytes)
        tmp_doc_path = tmp.name

    tmp_docx_path = tmp_doc_path + ".docx"
    modified_docx_path = tmp_doc_path + ".modified.docx"
    output_doc_path = tmp_doc_path + ".output.doc"

    try:
        result = subprocess.run(
            ["textutil", "-convert", "docx", "-output", tmp_docx_path, tmp_doc_path],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode != 0:
            raise RuntimeError(f"textutil doc->docx failed: {result.stderr}")

        with open(tmp_docx_path, "rb") as f:
            docx_bytes = f.read()

        modified_docx = apply_replacements_to_docx(docx_bytes, replacements)

        with open(modified_docx_path, "wb") as f:
            f.write(modified_docx)

        result = subprocess.run(
            ["textutil", "-convert", "doc", "-output", output_doc_path, modified_docx_path],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode != 0:
            raise RuntimeError(f"textutil docx->doc failed: {result.stderr}")

        with open(output_doc_path, "rb") as f:
            return f.read()
    finally:
        for path in [tmp_doc_path, tmp_docx_path, modified_docx_path, output_doc_path]:
            if os.path.exists(path):
                os.unlink(path)


def build_replacement_pairs(mapping_data: dict, reverse: bool = False) -> list[tuple[str, str]]:
    """
    Build a sorted list of (old_text, new_text) pairs from mapping data.

    Args:
        mapping_data: Full mapping dictionary with "mappings" key
        reverse: If False, entity->placeholder (anonymize).
                 If True, placeholder->entity (de-anonymize).

    Returns:
        List of (old, new) tuples sorted by old_text length descending
    """
    pairs = {}
    for placeholder, info in mapping_data.get("mappings", {}).items():
        value = info.get("value", "")
        if reverse:
            pairs[placeholder] = value
        else:
            if value:
                pairs[value] = placeholder
            for alias in info.get("aliases", []):
                if alias:
                    pairs[alias] = placeholder

    return sorted(pairs.items(), key=lambda x: len(x[0]), reverse=True)


def save_mapping(mapping_dict: dict, filename: str) -> str:
    """
    Save mapping table as a JSON file.

    Args:
        mapping_dict: Mapping dictionary
        filename: Base filename (without extension)

    Returns:
        Path to the saved file
    """
    os.makedirs(MAPPINGS_DIR, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    save_filename = f"{filename}_mapping_{timestamp}.json"
    filepath = os.path.join(MAPPINGS_DIR, save_filename)

    with open(filepath, "w", encoding="utf-8") as f:
        json.dump(mapping_dict, f, ensure_ascii=False, indent=2)

    return filepath


def load_mapping(uploaded_file) -> dict:
    """
    Load mapping table from an uploaded JSON file.

    Args:
        uploaded_file: Streamlit UploadedFile object

    Returns:
        Mapping dictionary
    """
    content = uploaded_file.read()
    text = content.decode("utf-8")
    return json.loads(text)
